#Importing all required Dependencies
import streamlit as st
import docx
from googletrans import Translator
from docx import Document
#Function to extract text from a docx file
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def show_supply_side_page(): 
    st.title("\n")
    st.title("EduReka")
    st.subheader("- Notes Rendering Service for Creators/Educators")
    st.title("")
    #Uploading and saving file
    uploaded_file = st.file_uploader("Upload the notes in English(.docx):")
    if uploaded_file is not None:
        bytes_data = uploaded_file.getvalue() 
        with open(uploaded_file.name, 'wb') as f: #Saving the file locally for translation purpose
            f.write(uploaded_file.getbuffer())
        text = getText(str(uploaded_file.name))
        
        option = st.selectbox('Select Language you want your notes to be in:',('English', 'Hindi', 'Arabic', 'Bengali', 'Greek', 'Japanese', 'Tamil', 'Telugu'))
        
        #Dictionary Created to store language codes of few languages 
        lang_codes = {'English' : 'en', 'Hindi' : 'hi', 'Arabic':'ar', 'Bengali':'bn', 'Greek':'el', 'Japanese': 'ja', 'Tamil': 'ta','Telugu': 'te'}



        #Translator defined and executed
        translator = Translator()
        print("Translating...")
        out = translator.translate(str(text), dest=str(lang_codes[option]))
        #Saving translated text in a docx file
        document = Document()
        document.add_heading('Notes in '+str(option)+'\n', level=1)
        document.add_paragraph(str(out.text))
        document.save('rendered_notes.docx')

        #Download Option for users 
        with open("rendered_notes.docx", "rb") as file:
            btn = st.download_button(
                    label="DOWNLOAD",
                    data=file,
                    file_name="rendered_notes.docx",
                    mime="text/docx"
                  )